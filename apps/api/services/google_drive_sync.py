"""Google Drive synchronization service.

Syncs file metadata from Google Drive into cloud_documents table.
Files are NOT copied (GDPR compliance / secret professionnel Art. 458 C.P. belge).
Only metadata + optional Qdrant index are stored.

Uses the Google Drive API v3 via raw HTTP requests with Bearer token.
"""

import logging
from datetime import datetime, timezone
from typing import Optional
from uuid import UUID

import httpx
from sqlalchemy import select
from sqlalchemy.dialects.postgresql import insert as pg_insert
from sqlalchemy.ext.asyncio import AsyncSession

from packages.db.models.cloud_document import CloudDocument
from packages.db.models.cloud_sync_job import CloudSyncJob
from apps.api.services.oauth_engine import get_oauth_engine

logger = logging.getLogger(__name__)

DRIVE_BASE_URL = "https://www.googleapis.com/drive/v3"
DRIVE_TIMEOUT = 30.0

# Google MIME type -> is it a Google native doc (editable online)?
GOOGLE_NATIVE_TYPES = {
    "application/vnd.google-apps.document",
    "application/vnd.google-apps.spreadsheet",
    "application/vnd.google-apps.presentation",
    "application/vnd.google-apps.form",
    "application/vnd.google-apps.drawing",
}

GOOGLE_FOLDER_TYPE = "application/vnd.google-apps.folder"


def _get_edit_url(file_id: str, mime_type: str) -> Optional[str]:
    """Return the online edit URL for a Google file."""
    if mime_type == "application/vnd.google-apps.document":
        return f"https://docs.google.com/document/d/{file_id}/edit"
    elif mime_type == "application/vnd.google-apps.spreadsheet":
        return f"https://docs.google.com/spreadsheets/d/{file_id}/edit"
    elif mime_type == "application/vnd.google-apps.presentation":
        return f"https://docs.google.com/presentation/d/{file_id}/edit"
    elif mime_type == "application/vnd.google-apps.form":
        return f"https://docs.google.com/forms/d/{file_id}/edit"
    return None


class GoogleDriveSync:
    """Synchronizes Google Drive files into cloud_documents."""

    async def list_files(
        self,
        session: AsyncSession,
        token_id: UUID,
        tenant_id: UUID,
        folder_id: str = "root",
        page_token: Optional[str] = None,
    ) -> dict:
        """List files in a Drive folder.

        Returns: {files: [...], nextPageToken: str | None}
        """
        oauth_engine = get_oauth_engine()
        access_token = await oauth_engine.get_valid_token(session, token_id)

        params = {
            "q": f"'{folder_id}' in parents and trashed = false",
            "fields": "nextPageToken,files(id,name,mimeType,size,webViewLink,thumbnailLink,modifiedTime,lastModifyingUser,parents,md5Checksum)",
            "pageSize": 100,
            "orderBy": "modifiedTime desc",
        }
        if page_token:
            params["pageToken"] = page_token

        async with httpx.AsyncClient(timeout=DRIVE_TIMEOUT) as client:
            response = await client.get(
                f"{DRIVE_BASE_URL}/files",
                headers={"Authorization": f"Bearer {access_token}"},
                params=params,
            )
            response.raise_for_status()
            return response.json()
    async def sync_folder(
        self,
        session: AsyncSession,
        token_id: UUID,
        tenant_id: UUID,
        folder_id: str = "root",
        recursive: bool = True,
        job_id: Optional[UUID] = None,
        parent_path: str = "",
    ) -> dict:
        """Recursively sync a Google Drive folder.

        Returns: {synced: int, errors: int}
        """
        synced = 0
        errors = 0
        page_token = None

        while True:
            try:
                result = await self.list_files(
                    session, token_id, tenant_id, folder_id, page_token
                )
            except Exception as e:
                logger.error(f"Failed to list Drive folder {folder_id}: {e}")
                errors += 1
                break

            files = result.get("files", [])

            for file_data in files:
                try:
                    await self._upsert_document(
                        session, token_id, tenant_id, file_data, parent_path
                    )
                    synced += 1

                    # Update job progress
                    if job_id:
                        await self._update_job_progress(session, job_id, synced)

                    # Recurse into subfolders
                    if recursive and file_data.get("mimeType") == GOOGLE_FOLDER_TYPE:
                        subfolder_path = f"{parent_path}/{file_data['name']}".lstrip("/")
                        sub = await self.sync_folder(
                            session, token_id, tenant_id,
                            file_data["id"], recursive, job_id, subfolder_path
                        )
                        synced += sub["synced"]
                        errors += sub["errors"]

                except Exception as e:
                    logger.error(f"Failed to sync Drive file {file_data.get('id')}: {e}")
                    errors += 1

            page_token = result.get("nextPageToken")
            if not page_token:
                break

        return {"synced": synced, "errors": errors}

    async def incremental_sync(
        self,
        session: AsyncSession,
        token_id: UUID,
        tenant_id: UUID,
        since: Optional[datetime] = None,
    ) -> dict:
        """Sync only files modified since a given datetime."""
        oauth_engine = get_oauth_engine()
        access_token = await oauth_engine.get_valid_token(session, token_id)

        # Build query
        if since:
            since_rfc3339 = since.strftime("%Y-%m-%dT%H:%M:%S")
            q = f"modifiedTime > '{since_rfc3339}' and trashed = false"
        else:
            q = "trashed = false"

        params = {
            "q": q,
            "fields": "nextPageToken,files(id,name,mimeType,size,webViewLink,thumbnailLink,modifiedTime,lastModifyingUser,parents,md5Checksum)",
            "pageSize": 100,
            "orderBy": "modifiedTime desc",
        }

        synced = 0
        errors = 0
        page_token = None

        while True:
            async with httpx.AsyncClient(timeout=DRIVE_TIMEOUT) as client:
                response = await client.get(
                    f"{DRIVE_BASE_URL}/files",
                    headers={"Authorization": f"Bearer {access_token}"},
                    params={**params, **({"-pageToken": page_token} if page_token else {})},
                )
                response.raise_for_status()
                result = response.json()

            for file_data in result.get("files", []):
                try:
                    await self._upsert_document(session, token_id, tenant_id, file_data, "")
                    synced += 1
                except Exception as e:
                    logger.error(f"Incremental sync error for {file_data.get('id')}: {e}")
                    errors += 1

            page_token = result.get("nextPageToken")
            if not page_token:
                break

        return {"synced": synced, "errors": errors}
    async def search_files(
        self,
        session: AsyncSession,
        token_id: UUID,
        tenant_id: UUID,
        query: str,
    ) -> list[dict]:
        """Search Drive files by full-text query."""
        oauth_engine = get_oauth_engine()
        access_token = await oauth_engine.get_valid_token(session, token_id)

        params = {
            "q": f"fullText contains '{query.replace(chr(39), chr(39)*2)}' and trashed = false",
            "fields": "files(id,name,mimeType,size,webViewLink,modifiedTime)",
            "pageSize": 20,
        }

        async with httpx.AsyncClient(timeout=DRIVE_TIMEOUT) as client:
            response = await client.get(
                f"{DRIVE_BASE_URL}/files",
                headers={"Authorization": f"Bearer {access_token}"},
                params=params,
            )
            response.raise_for_status()
            return response.json().get("files", [])

    async def get_file_content(
        self,
        session: AsyncSession,
        token_id: UUID,
        file_id: str,
        mime_type: str,
    ) -> bytes:
        """Download file content.

        For native Google formats, exports as PDF.
        """
        oauth_engine = get_oauth_engine()
        access_token = await oauth_engine.get_valid_token(session, token_id)

        async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
            if mime_type in GOOGLE_NATIVE_TYPES:
                # Export native Google format
                export_mime = "application/pdf"
                if mime_type == "application/vnd.google-apps.spreadsheet":
                    export_mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                
                response = await client.get(
                    f"{DRIVE_BASE_URL}/files/{file_id}/export",
                    headers={"Authorization": f"Bearer {access_token}"},
                    params={"mimeType": export_mime},
                )
            else:
                # Download binary file
                response = await client.get(
                    f"{DRIVE_BASE_URL}/files/{file_id}",
                    headers={"Authorization": f"Bearer {access_token}"},
                    params={"alt": "media"},
                )
            
            response.raise_for_status()
            return response.content

    async def get_file_text(
        self,
        session: AsyncSession,
        token_id: UUID,
        file_id: str,
        mime_type: str,
    ) -> str:
        """Extract text content for indexing."""
        oauth_engine = get_oauth_engine()
        access_token = await oauth_engine.get_valid_token(session, token_id)

        # For Google Docs, use Docs API to get structured content
        if mime_type == "application/vnd.google-apps.document":
            async with httpx.AsyncClient(timeout=DRIVE_TIMEOUT) as client:
                response = await client.get(
                    f"https://docs.googleapis.com/v1/documents/{file_id}",
                    headers={"Authorization": f"Bearer {access_token}"},
                )
                if response.status_code == 200:
                    doc = response.json()
                    return _extract_docs_text(doc)

        # For other types, export and extract
        try:
            content = await self.get_file_content(session, token_id, file_id, mime_type)
            return _extract_text_from_bytes(content, mime_type)
        except Exception as e:
            logger.warning(f"Could not extract text from {file_id}: {e}")
            return ""
    async def _upsert_document(
        self,
        session: AsyncSession,
        token_id: UUID,
        tenant_id: UUID,
        file_data: dict,
        parent_path: str,
    ) -> CloudDocument:
        """Insert or update a cloud_document record."""
        file_id = file_data["id"]
        mime_type = file_data.get("mimeType", "")
        is_folder = mime_type == GOOGLE_FOLDER_TYPE

        # Build URLs
        web_url = file_data.get("webViewLink")
        edit_url = _get_edit_url(file_id, mime_type) if not is_folder else None

        # Modified time
        modified_time = None
        if file_data.get("modifiedTime"):
            try:
                modified_time = datetime.fromisoformat(
                    file_data["modifiedTime"].replace("Z", "+00:00")
                )
            except Exception:
                pass

        # Last modifier
        last_modifier = None
        if file_data.get("lastModifyingUser"):
            last_modifier = file_data["lastModifyingUser"].get("displayName")

        # Path
        file_path = f"{parent_path}/{file_data['name']}".lstrip("/")

        # Check existing
        stmt = select(CloudDocument).where(
            CloudDocument.tenant_id == tenant_id,
            CloudDocument.provider == "google_drive",
            CloudDocument.external_id == file_id,
        )
        result = await session.execute(stmt)
        existing = result.scalar_one_or_none()

        if existing:
            # Update
            existing.name = file_data["name"]
            existing.mime_type = mime_type
            existing.size_bytes = int(file_data["size"]) if file_data.get("size") else None
            existing.web_url = web_url
            existing.edit_url = edit_url
            existing.thumbnail_url = file_data.get("thumbnailLink")
            existing.last_modified_at = modified_time
            existing.last_modified_by = last_modifier
            existing.content_hash = file_data.get("md5Checksum")
            existing.path = file_path
            existing.external_parent_id = (
                file_data["parents"][0] if file_data.get("parents") else None
            )
            await session.flush()
            return existing
        else:
            doc = CloudDocument(
                tenant_id=tenant_id,
                oauth_token_id=token_id,
                provider="google_drive",
                external_id=file_id,
                external_parent_id=(
                    file_data["parents"][0] if file_data.get("parents") else None
                ),
                name=file_data["name"],
                mime_type=mime_type,
                size_bytes=int(file_data["size"]) if file_data.get("size") else None,
                web_url=web_url,
                edit_url=edit_url,
                thumbnail_url=file_data.get("thumbnailLink"),
                is_folder=is_folder,
                path=file_path,
                last_modified_at=modified_time,
                last_modified_by=last_modifier,
                content_hash=file_data.get("md5Checksum"),
                is_indexed=False,
                index_status="pending",
            )
            session.add(doc)
            await session.flush()
            return doc

    async def _update_job_progress(
        self, session: AsyncSession, job_id: UUID, processed: int
    ) -> None:
        """Update sync job progress."""
        result = await session.execute(
            select(CloudSyncJob).where(CloudSyncJob.id == job_id)
        )
        job = result.scalar_one_or_none()
        if job:
            job.processed_items = processed
            await session.flush()

def _extract_docs_text(doc: dict) -> str:
    """Extract plain text from Google Docs API response."""
    texts = []
    body = doc.get("body", {})
    
    def extract_from_content(content: list) -> None:
        for element in content:
            if "paragraph" in element:
                for para_element in element["paragraph"].get("elements", []):
                    if "textRun" in para_element:
                        texts.append(para_element["textRun"].get("content", ""))
            elif "table" in element:
                for row in element["table"].get("tableRows", []):
                    for cell in row.get("tableCells", []):
                        extract_from_content(cell.get("content", []))
    
    extract_from_content(body.get("content", []))
    return "".join(texts)


def _extract_text_from_bytes(content: bytes, mime_type: str) -> str:
    """Extract text from file bytes based on MIME type."""
    try:
        if "text/" in mime_type:
            return content.decode("utf-8", errors="ignore")
        elif mime_type == "application/pdf":
            try:
                import pdfplumber
                import io
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    return "
".join(
                        page.extract_text() or "" for page in pdf.pages
                    )
            except ImportError:
                return ""
        elif "wordprocessingml" in mime_type or mime_type.endswith(".docx"):
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(content))
                return "
".join(para.text for para in doc.paragraphs)
            except ImportError:
                return ""
    except Exception:
        pass
    return ""


# Singleton
_google_drive_sync: GoogleDriveSync | None = None


def get_google_drive_sync() -> GoogleDriveSync:
    global _google_drive_sync
    if _google_drive_sync is None:
        _google_drive_sync = GoogleDriveSync()
    return _google_drive_sync
